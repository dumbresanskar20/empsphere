from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    return p

doc = Document()

# Title
title = doc.add_heading('Employee Management System Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Introduction
add_heading(doc, '1. Introduction', 1)
add_paragraph(doc, "The Employee Management System (EMS) is a comprehensive web-based application designed to streamline and automate the human resource operations within an organization. It provides a centralized platform for managing employee details, department assignments, leave requests, payroll processing, and document verification. The system features a robust role-based access control mechanism, distinguishing between Administrators, who have full oversight and management capabilities, and Employees, who can manage their personal profiles, submit leave requests, and view their salary slips. By transitioning from manual to digital processes, the EMS ensures higher efficiency, accuracy, and security of sensitive organizational data.")

# 2. Problem Statement
add_heading(doc, '2. Problem Statement', 1)
add_paragraph(doc, "Traditional methods of managing employee records rely heavily on paper-based systems or disconnected spreadsheets. These methods present several significant challenges:\n"
"- Data Inconsistency: Maintaining multiple records across different departments often leads to mismatched and outdated information.\n"
"- Time-Consuming Processes: Manual leave processing and payroll calculations are tedious and prone to human error.\n"
"- Security Risks: Physical files and unprotected digital spreadsheets are vulnerable to unauthorized access and loss.\n"
"- Lack of Transparency: Employees often lack real-time access to their leave balances, salary details, and approval statuses.\n"
"To address these inefficiencies, a centralized, automated, and secure system is essential.")

# 3. Proposed System
add_heading(doc, '3. Proposed System', 1)
add_paragraph(doc, "The proposed Employee Management System is a Full-Stack web application tailored to resolve the issues inherent in traditional HR management. Key features of the proposed system include:\n"
"- Secure Authentication: JWT-based secure login with role-based access for Admins and Employees.\n"
"- Centralized Database: Utilizing MongoDB to store all organizational data in a unified, accessible, and secure format.\n"
"- Leave Management: Automated workflow for employees to request leave and for admins to approve or reject them.\n"
"- Department Management: Structured categorization of employees into specific organizational units.\n"
"- Analytics Dashboard: Real-time visual insights (using Chart.js) into employee distribution, leave statuses, and departmental statistics.\n"
"- Document Uploads: Secure handling of employee verification documents using cloud or local storage mechanisms.")

# 4. System Requirements
add_heading(doc, '4. System Requirements', 1)
add_heading(doc, '4.1 Hardware Requirements', 2)
add_paragraph(doc, "- Processor: Intel Core i3 or equivalent (minimum)\n"
"- RAM: 4 GB (8 GB recommended)\n"
"- Storage: 10 GB of free space for application deployment and basic data storage\n"
"- Network: Stable internet connection for database connectivity and client access")

add_heading(doc, '4.2 Software Requirements', 2)
add_paragraph(doc, "- Operating System: Windows, macOS, or Linux\n"
"- Frontend: HTML5, CSS3, JavaScript, Bootstrap/TailwindCSS (optional), Chart.js\n"
"- Backend Runtime: Node.js (v14.x or higher)\n"
"- Backend Framework: Express.js\n"
"- Database: MongoDB (Atlas Cloud or Local Server)\n"
"- Version Control: Git\n"
"- Web Browser: Google Chrome, Mozilla Firefox, or Safari (latest versions)")

# 5. System Design
add_heading(doc, '5. System Design', 1)
add_heading(doc, '5.1 System Architecture', 2)
add_paragraph(doc, "The system follows a modern Client-Server Architecture based on the MVC (Model-View-Controller) design pattern:\n"
"- Client (Frontend): Developed using HTML, CSS, and Vanilla JavaScript. It handles the user interface, form validations, and asynchronous API calls (AJAX/Fetch) to communicate with the server.\n"
"- Server (Backend): Built with Node.js and Express.js, it provides RESTful API endpoints, handles business logic, authenticates users (JWT), and processes data.\n"
"- Database (Model): MongoDB serves as the NoSQL database, storing unstructured and semi-structured data securely.")

add_heading(doc, '5.2 ER Diagram', 2)
add_paragraph(doc, "Entities and Relationships:\n"
"- User (Employee/Admin): Central entity. Has relationships with Department, Leaves, and Salary.\n"
"- Department: One-to-Many relationship with Users (One department has many employees).\n"
"- Leave: Many-to-One relationship with User (An employee can have multiple leave requests).\n"
"- Salary: Many-to-One relationship with User (An employee receives multiple salary records over time).")

add_heading(doc, '5.3 Schema', 2)
add_paragraph(doc, "The database utilizes Document-oriented schemas:\n"
"- User Schema: _id, employeeId, name, email, password (hashed), role, department, designation, salary, status.\n"
"- Department Schema: _id, name, description, createdAt.\n"
"- Leave Schema: _id, userId (ref: User), type, startDate, endDate, reason, status (Pending/Approved/Rejected).\n"
"- Document Schema: _id, userId (ref: User), fileName, fileUrl, uploadDate.")

# 6. Database Design
add_heading(doc, '6. Database Design', 1)
add_heading(doc, '6.1 Tables (Collections)', 2)
add_paragraph(doc, "1. Users: Stores credentials and profile information.\n"
"2. Departments: Stores structural units of the organization.\n"
"3. Leaves: Stores history of time-off requests.\n"
"4. Salaries: Logs of monthly or periodic payments.")

add_heading(doc, '6.2 Primary/Foreign Keys', 2)
add_paragraph(doc, "In MongoDB, relationships are maintained using references:\n"
"- Primary Key: Every document automatically receives a unique '_id' field.\n"
"- Foreign Key Equivalent: Fields like 'userId' in the Leaves collection store the Object ID of a document in the Users collection, linking the records.")

add_heading(doc, '6.3 Normalization (1NF, 2NF, 3NF)', 2)
add_paragraph(doc, "Although MongoDB is a NoSQL database and often denormalized for performance, the logical design principles of normalization apply to structure our data cleanly.")

# Normalization Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Normal Form'
hdr_cells[1].text = 'Description / Rule'
hdr_cells[2].text = 'Application in EMS'

row_cells = table.add_row().cells
row_cells[0].text = 'First Normal Form (1NF)'
row_cells[1].text = 'Ensure each column contains atomic values, and each record is unique.'
row_cells[2].text = 'An Employee collection separates first name, last name, and contact array. No multiple comma-separated values in a single string field.'

row_cells = table.add_row().cells
row_cells[0].text = 'Second Normal Form (2NF)'
row_cells[1].text = 'Meets 1NF and all non-key attributes are fully functional dependent on the primary key.'
row_cells[2].text = 'Department details (like Department Location) are moved to a separate "Departments" collection instead of duplicating them in every Employee record.'

row_cells = table.add_row().cells
row_cells[0].text = 'Third Normal Form (3NF)'
row_cells[1].text = 'Meets 2NF and there are no transitive dependencies (non-key attributes depending on other non-key attributes).'
row_cells[2].text = 'Calculated fields like "Net Salary" are either computed dynamically or stored carefully. Salary scales based on Designation are stored in a linked table rather than redundantly in the user document.'

# 7. Implementation
add_heading(doc, '7. Implementation', 1)
add_heading(doc, '7.1 Source Code Overview', 2)
add_paragraph(doc, "The implementation is divided into two main directories:\n"
"- /frontend: Contains HTML pages (index.html, admin.html, employee.html), CSS styling for responsive design, and vanilla JavaScript for DOM manipulation and API integration.\n"
"- /backend: Contains the Express server setup (server.js), Mongoose models (User.js, Leave.js, Department.js), and route controllers (authRoutes.js, adminRoutes.js).")

add_heading(doc, '7.2 Screenshots', 2)
add_paragraph(doc, "[Insert Screenshot of Login Page Here]\n"
"[Insert Screenshot of Admin Dashboard with Charts Here]\n"
"[Insert Screenshot of Employee Profile Here]\n"
"(Note: Attach actual screenshots of the responsive UI in the final printed copy.)")

# 8. Testing
add_heading(doc, '8. Testing', 1)
add_paragraph(doc, "- Unit Testing: Individual API endpoints (e.g., POST /api/auth/login) were tested using tools like Postman to ensure proper request validation and response handling.\n"
"- Integration Testing: Verified that the frontend successfully communicates with the backend, such as ensuring leave requests submitted from the client accurately update the MongoDB database.\n"
"- UI/UX Testing: The responsive design was tested across various screen sizes (Mobile, Tablet, Desktop) to ensure the hamburger menu and fluid layouts perform correctly.\n"
"- Security Testing: Verified that unauthorized users cannot access Admin pages and that JWT tokens expire and validate correctly.")

# 9. Advantages
add_heading(doc, '9. Advantages', 1)
add_paragraph(doc, "1. Efficiency: Automates routine HR tasks, saving significant time.\n"
"2. Accessibility: Web-based system allows access from anywhere with an internet connection.\n"
"3. Security: Encrypted passwords and JWT authentication protect sensitive data.\n"
"4. Scalability: Node.js and MongoDB easily handle an increasing number of employee records.\n"
"5. Insights: The dynamic dashboard provides immediate visual feedback on company statistics.")

# 10. Limitations
add_heading(doc, '10. Limitations', 1)
add_paragraph(doc, "1. Internet Dependency: Requires continuous internet/network access to function.\n"
"2. Learning Curve: Staff may require brief training to transition from older manual systems.\n"
"3. Deployment Overhead: Requires server hosting and database maintenance for production deployment.")

# 11. Conclusion
add_heading(doc, '11. Conclusion', 1)
add_paragraph(doc, "The Employee Management System successfully digitizes HR operations, providing a robust, scalable, and user-friendly platform. By leveraging modern web technologies (Node.js, Express, MongoDB, and responsive frontend design), the system not only solves the inefficiencies of manual tracking but also lays a strong foundation for future technological integrations within the organization.")

# 12. Future Scope
add_heading(doc, '12. Future Scope', 1)
add_paragraph(doc, "- Mobile Application: Developing native iOS and Android applications for greater accessibility.\n"
"- Attendance Tracking: Integration with biometric or IP-based attendance tracking systems.\n"
"- AI Analytics: Implementing predictive analytics for employee turnover and performance metrics.\n"
"- Advanced Payroll System: Integration with regional tax calculation APIs and bank auto-deposits.")

# 13. Reference
add_heading(doc, '13. Reference', 1)
add_paragraph(doc, "- Node.js Documentation: https://nodejs.org/docs\n"
"- Express.js Documentation: https://expressjs.com/\n"
"- MongoDB Atlas & Mongoose: https://mongoosejs.com/\n"
"- Chart.js Documentation: https://www.chartjs.org/\n"
"- JSON Web Tokens: https://jwt.io/")

doc.save('Employee_Management_System_Documentation.docx')
print("Document saved successfully.")
