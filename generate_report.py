import os
import subprocess
import sys

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Employee Management System Project Report', 0)

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('The Employee Management System (EmpSphere) is a full-stack web application designed to streamline human resource processes within an organization. It provides a centralized platform for managing employee records, tracking leaves, monitoring salaries, and handling document approvals. The system caters to both Administrators and Employees with dedicated portals and role-based access controls, ensuring data security and efficient operational management.')

# 2. Problem Statement
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph('Traditional employee management involves manual record-keeping, disparate systems, and time-consuming administrative tasks. Organizations often struggle with tracking employee leaves, processing salaries accurately, and managing employee documents securely. The lack of an integrated system leads to inefficiencies, data inconsistencies, and delayed decision-making processes.')

# 3. Proposed System
doc.add_heading('3. Proposed System', level=1)
doc.add_paragraph('The proposed system is a robust, web-based Employee Management System that automates and integrates core HR functions. Key features include:')
doc.add_paragraph('Admin Dashboard: Centralized view of analytics (using Chart.js), employee approval workflows, leave tracking, department management, and salary administration.', style='List Bullet')
doc.add_paragraph('Employee Portal: Self-service portal for employees to apply for leaves, upload required documents, manage their profiles, and view personal analytics.', style='List Bullet')
doc.add_paragraph('Role-Based Authentication: Secure JWT-based authentication with bcrypt password hashing for data protection.', style='List Bullet')
doc.add_paragraph('Document Management: Secure document upload and storage system with administrative approval mechanisms.', style='List Bullet')

# 4. System Requirements
doc.add_heading('4. System Requirements', level=1)
doc.add_heading('Hardware Requirements', level=2)
doc.add_paragraph('Processor: Intel Core i3 or higher / Equivalent AMD Processor', style='List Bullet')
doc.add_paragraph('RAM: 4 GB minimum (8 GB recommended)', style='List Bullet')
doc.add_paragraph('Storage: 10 GB of free disk space', style='List Bullet')
doc.add_paragraph('Network: Stable Internet connection', style='List Bullet')

doc.add_heading('Software Requirements', level=2)
doc.add_paragraph('Operating System: Windows 10/11, macOS, or Linux', style='List Bullet')
doc.add_paragraph('Frontend: HTML5, CSS3, JavaScript (Vanilla JS/React.js), Chart.js', style='List Bullet')
doc.add_paragraph('Backend: Node.js, Express.js', style='List Bullet')
doc.add_paragraph('Database: MongoDB Atlas (NoSQL)', style='List Bullet')
doc.add_paragraph('Tools: VS Code, Web Browser, Postman', style='List Bullet')

# 5. System Design
doc.add_heading('5. System Design', level=1)
doc.add_heading('System Architecture', level=2)
doc.add_paragraph('The system follows a client-server (3-tier) architecture:')
doc.add_paragraph('Presentation Layer: The user interface built with HTML, CSS, and JS. It communicates with the backend via RESTful APIs.', style='List Bullet')
doc.add_paragraph('Application Layer: Node.js and Express.js backend handling business logic, authentication (JWT), file uploads (Multer), and database interactions.', style='List Bullet')
doc.add_paragraph('Data Layer: MongoDB Atlas storing users, leaves, salaries, documents, and departments.', style='List Bullet')

doc.add_heading('ER Diagram & Schema', level=2)
doc.add_paragraph('The schema is designed using Mongoose for MongoDB. Key collections include Users, Leaves, Salaries, Departments, and Documents. Although MongoDB is NoSQL, the logical relationship revolves around the User (Employee/Admin). A User has multiple Leaves, Salaries, and Documents associated with their Employee ID.')

# 6. Database Design
doc.add_heading('6. Database Design', level=1)
doc.add_heading('Tables (Collections)', level=2)
doc.add_paragraph('Users: Stores employee details, role, status, credentials, and embedded document metadata.', style='List Bullet')
doc.add_paragraph('Leaves: Stores leave requests, start/end dates, reason, and status (Pending/Approved/Rejected).', style='List Bullet')
doc.add_paragraph('Salaries: Stores salary records, deductions, bonuses, and payment status for employees.', style='List Bullet')
doc.add_paragraph('Departments: Stores department names and statistics.', style='List Bullet')

doc.add_heading('Primary/Foreign Keys (Logical Relationships)', level=2)
doc.add_paragraph('Primary Key: _id (ObjectId) automatically generated by MongoDB for each document.', style='List Bullet')
doc.add_paragraph('Foreign Key Equivalent: The system uses employeeId and userId references across collections (Leaves, Salaries) to link data back to the Users collection.', style='List Bullet')

doc.add_heading('Normalization (1NF, 2NF, 3NF)', level=2)
doc.add_paragraph('Since MongoDB is a document-based NoSQL database, traditional normalization rules are applied differently. However, the schema logical design follows basic principles:')
doc.add_paragraph('1NF: Each field contains atomic values (e.g., strings, numbers, dates).', style='List Bullet')
doc.add_paragraph('2NF: Partial dependencies are eliminated by separating Leaves and Salaries into their own collections rather than deeply embedding everything in the User document.', style='List Bullet')
doc.add_paragraph('3NF: Transitive dependencies are minimized by storing reference IDs (userId) and fetching details via population or application-level joins.', style='List Bullet')

# 7. Implementation
doc.add_heading('7. Implementation', level=1)
doc.add_heading('Source Code', level=2)
doc.add_paragraph('The implementation uses a clear folder structure with "frontend" and "backend" directories. The backend utilizes Express routers for defining API endpoints (/api/auth, /api/users, /api/leaves). Mongoose models define the data structure. The frontend uses Fetch API to consume these REST endpoints and dynamically render data using DOM manipulation and Chart.js.')

doc.add_heading('Screenshots', level=2)
doc.add_paragraph('(Screenshots of Admin Dashboard, Employee Portal, Login/Signup, and File Upload features will be attached here manually by the developer)')

# 8. Testing
doc.add_heading('8. Testing', level=1)
doc.add_paragraph('Testing strategies applied to the system include:')
doc.add_paragraph('Unit Testing: Ensuring individual functions like password hashing (bcrypt) and JWT generation work correctly.', style='List Bullet')
doc.add_paragraph('Integration Testing: Testing API endpoints with Postman to ensure backend communicates with MongoDB correctly.', style='List Bullet')
doc.add_paragraph('UI/UX Testing: Ensuring the frontend is responsive across desktop and mobile devices.', style='List Bullet')
doc.add_paragraph('Security Testing: Validating JWT token expiration, testing protected routes, and ensuring unapproved users cannot access the system.', style='List Bullet')

# 9. Advantages
doc.add_heading('9. Advantages', level=1)
doc.add_paragraph('Centralized Data Management: All employee records are stored in a single secure location.', style='List Bullet')
doc.add_paragraph('Automated Workflows: Document approvals and leave requests are streamlined.', style='List Bullet')
doc.add_paragraph('Data Security: Secure authentication and password hashing protect sensitive information.', style='List Bullet')
doc.add_paragraph('Accessibility: Web-based interface allows access from any device with an internet connection.', style='List Bullet')
doc.add_paragraph('Real-time Analytics: Dynamic charts provide instant insights into employee demographics and department distribution.', style='List Bullet')

# 10. Limitations
doc.add_heading('10. Limitations', level=1)
doc.add_paragraph('Internet Dependency: Requires a stable internet connection as it relies on MongoDB Atlas and web interfaces.', style='List Bullet')
doc.add_paragraph('No Offline Mode: Cannot be used offline without internet connectivity.', style='List Bullet')
doc.add_paragraph('Limited Initial Modules: Does not currently include advanced modules like payroll generation (tax calculations) or performance appraisals.', style='List Bullet')

# 11. Conclusion
doc.add_heading('11. Conclusion', level=1)
doc.add_paragraph('The Employee Management System successfully addresses the challenges of traditional HR management by providing an integrated, secure, and user-friendly web application. By leveraging modern web technologies (Node.js, Express, MongoDB), the system delivers a scalable solution that improves administrative efficiency and empowers employees through a self-service portal.')

# 12. Future Scope
doc.add_heading('12. Future Scope', level=1)
doc.add_paragraph('Advanced Payroll Processing: Integrating tax calculations, automated payslip generation, and direct bank transfers.', style='List Bullet')
doc.add_paragraph('Performance Management: Adding modules for employee reviews, goals tracking, and appraisals.', style='List Bullet')
doc.add_paragraph('Mobile Application: Developing a dedicated mobile app (e.g., using React Native) for push notifications and easier access.', style='List Bullet')
doc.add_paragraph('AI Integration: Implementing chatbots for HR queries or predictive analytics for employee retention.', style='List Bullet')

# 13. Reference
doc.add_heading('13. Reference', level=1)
doc.add_paragraph('Node.js Documentation: https://nodejs.org/docs/', style='List Bullet')
doc.add_paragraph('Express.js Documentation: https://expressjs.com/', style='List Bullet')
doc.add_paragraph('MongoDB & Mongoose: https://mongoosejs.com/', style='List Bullet')
doc.add_paragraph('Chart.js Documentation: https://www.chartjs.org/', style='List Bullet')
doc.add_paragraph('MDN Web Docs (HTML/CSS/JS): https://developer.mozilla.org/', style='List Bullet')

doc.save('Employee_Management_System_Report.docx')
print('Report generated successfully as Employee_Management_System_Report.docx')
